import os
import io
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, RGBColor
from langchain.prompts import ChatPromptTemplate
from langchain.chains import LLMChain
from langchain_community.llms import Ollama
from langchain.llms import OpenAI

# Function to select LLM based on user input
def select_llm(api_key):
    if api_key:
        llm = OpenAI(temperature=0, model="gpt-4", openai_api_key=api_key)
    else:
        llm = Ollama(model="llama3.1")
    return llm

# Function to calculate ratios
def calculate_ratios(df):
    ratios = {}
    if 'IDV' in df.columns:
        ratios['idv_count'] = df['IDV'].count()
    if 'Sexo' in df.columns:
        sexo_counts = df['Sexo'].value_counts(normalize=True)
        ratios['sexo_proportion'] = sexo_counts.to_dict()
        ratios['sexo_counts'] = df['Sexo'].value_counts().to_dict()
    if 'Peso al nacer' in df.columns:
        peso_stats = df['Peso al nacer'].describe()
        ratios['peso_al_nacer_stats'] = peso_stats.to_dict()
    if 'Toro Padre' in df.columns:
        toro_proportion = df['Toro Padre'].value_counts(normalize=True)
        ratios['toro_proportion'] = toro_proportion.to_dict()
        ratios['toro_counts'] = df['Toro Padre'].value_counts().to_dict()
    total_cows = 195
    ratios['calving_progress'] = len(df) / total_cows
    if 'Recorredor' in df.columns and 'Fecha' in df.columns:
        recorredor_counts = df['Recorredor'].value_counts()
        ratios['recorredor_counts'] = recorredor_counts.to_dict()
        last_fecha = df.groupby('Recorredor')['Fecha'].max()
        ratios['recorredor_last_fecha'] = last_fecha.dt.strftime('%d-%m-%Y').to_dict()
    if 'Fecha' in df.columns:
        fixed_date = pd.to_datetime('2024-08-03')
        df['Fecha'] = pd.to_datetime(df['Fecha'], errors='coerce')
        df['calf_age_days'] = (fixed_date - df['Fecha']).dt.days
        ratios['calf_age_days_stats'] = df['calf_age_days'].describe().to_dict()
    if 'CC Parto' in df.columns:
        cc_parto_proportion = df['CC Parto'].value_counts(normalize=True)
        ratios['cc_parto_proportion'] = cc_parto_proportion.to_dict()
        ratios['cc_parto_counts'] = df['CC Parto'].value_counts().to_dict()
    return ratios

# Function to plot data and add to Word document
def plot_to_word(document, ratios, title, key):
    if key in ratios:
        fig, ax = plt.subplots(figsize=(8, 6))
        if 'counts' in key:
            ax.pie(ratios[key].values(), labels=ratios[key].keys(), autopct='%1.1f%%')
        else:
            ax.bar(ratios[key].keys(), ratios[key].values())
            ax.set_xticklabels(ratios[key].keys(), rotation=45, ha='right')
            for i, v in enumerate(ratios[key].values()):
                ax.text(i, v + 0.01 * max(ratios[key].values()), f"{v:.2f}", ha='center', va='bottom')
        ax.set_title(title)
        image_stream = io.BytesIO()
        plt.savefig(image_stream, format='png')
        plt.close(fig)
        image_stream.seek(0)
        document.add_picture(image_stream, width=Inches(6))
        image_stream.close()

# Function to generate comments using LLM
def generate_comments(llm, ratios):
    prompt_template = """
    You are an AI assistant with expertise in cattle management. Please provide detailed comments and translate to spanish on the following ratios and statistics:
    - IDV Count: {idv_count}
    - Proportion of Sexo (Macho and Hembra): {sexo_proportion}
    - Descriptive Statistics of Peso al Nacer: {peso_al_nacer_stats}
    - Proportion of Toro Padre: {toro_proportion}
    - Calving Progress: {calving_progress:.2%}
    - Recorredor Counts: {recorredor_counts}
    - Last Fecha per Recorredor: {recorredor_last_fecha}
    - Descriptive Statistics of Calf Age (Days): {calf_age_days_stats}
    - Finally make a couple of comment about bovine respiratory disease in early stages of the cattle life
    """
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = LLMChain(llm=llm, prompt=prompt)
    comments = chain.run(**ratios)
    return comments

# Function to generate summary report
def generate_summary_report(ratios, comments, logo_path):
    report = Document()
    style = report.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.color.rgb = RGBColor(0, 0, 0)

    table = report.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))

    cell = table.cell(0, 1)
    cell.text = 'Calving Summary Report'
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = 2400

    report.add_heading('Detailed Comments and Translations', level=1)

    comments_sections = comments.split('\n\n')
    for section in comments_sections:
        report.add_paragraph(section)
        if 'Proportion of Sexo' in section:
            plot_to_word(report, ratios, 'Proportion of Sexo (Macho and Hembra)', 'sexo_counts')
        elif 'Descriptive Statistics of Peso al Nacer' in section:
            plot_to_word(report, ratios, 'Descriptive Statistics of Peso al Nacer', 'peso_al_nacer_stats')
        elif 'Proportion of Toro Padre' in section:
            plot_to_word(report, ratios, 'Proportion of Toro Padre', 'toro_counts')
        elif 'Descriptive Statistics of Calf Age (Days)' in section:
            plot_to_word(report, ratios, 'Descriptive Statistics of Calf Age (Days)', 'calf_age_days_stats')
        elif 'Proportion of CC Parto' in section:
            plot_to_word(report, ratios, 'Proportion of CC Parto', 'cc_parto_counts')

    return report

# Load logo or handle missing logo
def load_logo(logo_path):
    if os.path.exists(logo_path):
        return logo_path
    else:
        st.error("Logo file not found.")
        return None

# Prepare and show the download button
def prepare_download_button(summary_report):
    buffer = io.BytesIO()
    summary_report.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="Download Report",
        data=buffer,
        file_name="calving_summary_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Streamlit app logic
def main():
    st.title("Calving Summary Report Generator")
    api_key = st.text_input("Enter OpenAI API Key (optional, for GPT-4)")
    uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])

    if uploaded_file:
        try:
            df = pd.read_excel(uploaded_file)
            df['Fecha'] = pd.to_datetime(df['Fecha'], format='%d-%m-%Y', errors='coerce')
        except Exception as e:
            st.error(f"Failed to process the Excel file: {e}")
            return

        llm = select_llm(api_key)
        ratios = calculate_ratios(df)
        comments = generate_comments(llm, ratios)

        st.write("Ratios:", ratios)
        st.write("Generated Comments:", comments)

        if st.button("Generate Report"):
            logo_path = load_logo("logo.png")
            if logo_path:
                summary_report = generate_summary_report(ratios, comments, logo_path)
                prepare_download_button(summary_report)

if __name__ == "__main__":
    main()

